# -*- coding: utf-8 -*-
import os, re
from docx import Document
L = 13
p = r"D:\英语教学\邓兴华\第13课时\第13课时_配套练习_中等.docx"
doc = Document(p)
texts = []
for para in doc.paragraphs:
    if para.text.strip():
        texts.append(("P", para.text))
for ti, tbl in enumerate(doc.tables):
    for r in tbl.rows:
        for c in r.cells:
            t = c.text.strip()
            if t:
                texts.append(("T%s" % ti, t))
full = "\n".join(t for _, t in texts)
ids = sorted(set(re.findall(r"DXH\d+_L\d+", full)))
print("IDs:", len(ids), ids)
print("trace lines:")
for k, t in texts:
    if "溯源" in t or "DXH" in t:
        print("  [%s] %s" % (k, t[:100]))
print("para nums:")
nums = []
for k, t in texts:
    if k.startswith("P"):
        m = re.match(r"^\s*(\d{1,2})[.、．]\s*", t)
        if m:
            nums.append(int(m.group(1)))
print(sorted(set(nums)))
print("context 11-22:")
for k, t in texts:
    m = re.match(r"^\s*(\d{1,2})[.、．]\s*", t) if k.startswith("P") else None
    n = int(m.group(1)) if m else None
    if n and 11 <= n <= 22:
        print("  [%s] %s" % (k, t[:90]))