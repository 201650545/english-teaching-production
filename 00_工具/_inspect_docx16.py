# -*- coding: utf-8 -*-
"""提取 DOCX 段落结构，打印段落首行（带格式标记），用于对比练习卷结构。"""
import sys
from docx import Document

def dump(path, max_lines=140):
    print("=" * 70)
    print("FILE:", path)
    print("=" * 70)
    doc = Document(path)
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        sz = None
        bold = False
        if p.runs:
            sz = p.runs[0].font.size.pt if p.runs[0].font.size else None
            bold = p.runs[0].font.bold
        mark = "B" if bold else ""
        print("[%s%s] %s" % ("%s" % (round(sz,1) if sz else "?"), mark, t))
        n += 1
        if n >= max_lines:
            print("... (truncated)")
            break
    print("--- tables:", len(doc.tables))

if __name__ == "__main__":
    for f in sys.argv[1:]:
        dump(f)