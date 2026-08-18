# -*- coding: utf-8 -*-
from docx import Document
import re
for lesson, path in [(13, r"D:\英语教学\邓兴华\第13课时\第13课时_配套练习_中等.docx"),
                     (14, r"D:\英语教学\邓兴华\第14课时\第14课时_配套练习_中等.docx"),
                     (15, r"D:\英语教学\邓兴华\第15课时\第15课时_配套练习_中等.docx")]:
    print("="*40, f"L{lesson}", "="*40)
    doc = Document(path)
    print("表格数:", len(doc.tables))
    for ti, tab in enumerate(doc.tables):
        if ti >= 3: 
            print(f"  [表{ti}] (略)")
            continue
        print(f"  [表{ti}] rows={len(tab.rows)} cols={len(tab.columns)}")
        for r in tab.rows[:6]:
            cells = [c.text.strip()[:20] for c in r.cells]
            print("   |", " | ".join(cells))
    # search source id in all text incl tables
    alltxt = "\n".join(p.text for p in doc.paragraphs)
    for tab in doc.tables:
        for r in tab.rows:
            for c in r.cells:
                alltxt += "\n" + c.text
    ids = re.findall(r'DXH2026_L\d+[_\w]*', alltxt)
    print("溯源ID命中:", list(dict.fromkeys(ids))[:20] if ids else "无")
    print()