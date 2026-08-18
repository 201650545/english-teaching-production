# -*- coding: utf-8 -*-
import os
from docx import Document
ROOT = r"D:\英语教学\邓兴华"
for l in [21,22,23,24,25]:
    p = os.path.join(ROOT, "第%02d课时" % l, "第%02d课时_配套练习_中等.docx" % l)
    d = Document(p)
    paras = [x.text for x in d.paragraphs if x.text.strip()]
    print("="*60)
    print("L%d 段落数=%d" % (l, len(paras)))
    for t in paras[:10]:
        print("  H:", t[:60])
    for t in paras:
        if t.startswith("第") and ("部分" in t or "参考答案" in t or "答题卡" in t or "测试卷" in t or "配套" in t):
            print("  SEC:", t[:40])
    print("  表格数=%d" % len(d.tables))