# -*- coding: utf-8 -*-
"""邓兴华 L10-L15 配套练习 DOCX 3+2 同步（训练型）
规则：每篇阅读把低开放度细节题转无选项填空（下划线+（填空）），高开放度（主旨/词义/推断）保留有选项。
处理两种选项排版：合并式"A. X  B. Y  C. Z"单段 / 分行式 A/B/C 各自成段。
L10 课件契约仅 A/B（无 C 篇），配套 C 篇保留原样。
答案区：fill 题输出规范答案，opt 题保留原选项字母。
"""
import docx, os, re

base = r"d:\英语教学\邓兴华"
fill_answer = {
 "L10": {"1":"Chinese","2":"Tom","4":"To the library","5":"Miss Green"},
 "L11": {"1":"School uniforms","2":"Keep quiet","4":"Early","5":"The school rules",
         "8":"At eight o'clock","9":"Run"},
 "L12": {"1":"¥18","2":"¥18","4":"Because it's the writer's birthday",
         "5":"Vegetable salad and green tea","8":"Before a vowel sound","10":"In polite requests"},
 "L13": {"1":"To buy food for a special dinner","2":"Twenty yuan","3":"Sixty yuan",
         "6":"Her mother","7":"Five yuan","9":"Forty-five yuan",
         "12":"Because it helps you save money","13":"Check your change","15":"Compare prices at different stores"},
 "L14": {"1":"She is tall and thin","2":"Black","3":"She makes the class fun",
         "6":"Mr. Wang","7":"He teaches in a funny way","8":"He is tall and strong",
         "12":"From their parents","13":"Straight or curly","15":"Be kind to everyone"},
 "L15": {"1":"Sunny and warm","2":"Winter","3":"Could you please tell us the weather on Saturday?",
         "6":"The beach","7":"Hot and sunny","8":"Could I take my new sunglasses?",
         "11":"Four","12":"Spring","13":"Because they have a long vacation"},
}

jobs = [
 ("L10","第10课时","第10课时_配套练习_中等.docx"),
 ("L11","第11课时","第11课时_配套练习_中等_班规版.docx"),
 ("L12","第12课时","第12课时_配套练习_中等.docx"),
 ("L13","第13课时","第13课时_配套练习_中等.docx"),
 ("L14","第14课时","第14课时_配套练习_中等.docx"),
 ("L15","第15课时","第15课时_配套练习_中等.docx"),
]

def set_text(para, text):
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    para.add_run(text)

def remove_para(para):
    para._element.getparent().remove(para._element)

for lesson, sub, fn in jobs:
    src = os.path.join(base, sub, fn)
    if not os.path.exists(src):
        print(f"[{lesson}] MISSING {src}"); continue
    doc = docx.Document(src)
    paras = doc.paragraphs
    fa = fill_answer[lesson]
    converted = 0
    i = 0
    while i < len(paras):
        pa = paras[i]
        m = re.match(r'^(\d{1,2})\.\s', pa.text.strip())
        if m and m.group(1) in fa:
            # locate option paragraphs starting at i+1
            j = i+1
            while j < len(paras) and paras[j].text.strip()=="":
                j += 1
            if j < len(paras) and re.match(r'^A\.\s', paras[j].text.strip()):
                ablock = paras[j].text
                if re.search(r'\sB\.\s', ablock):
                    # combined format: single paragraph holds A/B/C
                    set_text(paras[j], "____________（填空）")
                else:
                    # split format: A paragraph becomes the blank; remove B/C paragraphs
                    set_text(paras[j], "____________（填空）")
                    k = j + 1
                    while k < len(paras):
                        tk = paras[k].text.strip()
                        if re.match(r'^[BC]\.\s', tk):
                            remove_para(paras[k])
                            k += 1
                        else:
                            break
                converted += 1
        i += 1

    # answer key update
    ans_updated = 0
    for pa in paras:
        t = pa.text.strip()
        m = re.match(r'^(\d{1,2})~(\d{1,2})\s+([ABC]+)$', t)
        if not m or int(m.group(1)) > 15:
            continue
        start, end = int(m.group(1)), int(m.group(2))
        letters = m.group(3)
        if len(letters) != (end-start+1):
            continue
        parts = []
        for k, qn in enumerate(range(start, end+1)):
            qs = str(qn)
            parts.append(f"{qn}. {fa[qs]}" if qs in fa else f"{qn}. {letters[k]}")
        set_text(pa, "　　".join(parts))
        ans_updated += 1

    stem = fn.replace(".docx","")
    out = os.path.join(base, sub, stem + "_3+2.docx")
    doc.save(out)
    print(f"[{lesson}] converted={converted} ans_updated={ans_updated} -> {os.path.basename(out)}")

print("DONE")