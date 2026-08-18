# -*- coding: utf-8 -*-
"""Final verification: answer distribution and compliance check."""
import re
from collections import Counter
from docx import Document

path = r"D:\英语教学\邓兴华\第08课时\第08课时_配套练习.docx"
doc = Document(path)

# Extract answer key
answers = {}
in_answer = False
for p in doc.paragraphs:
    text = p.text.strip()
    if "参考答案" in text:
        in_answer = True
        continue
    if in_answer:
        # Match patterns like "1~5 CABCA" or "11. C" or "26. beach"
        m = re.match(r"(\d+)~(\d+)\s+([A-Z]+)", text)
        if m:
            start, end, letters = int(m.group(1)), int(m.group(2)), m.group(3)
            for i, l in enumerate(letters):
                answers[start + i] = l
        else:
            m2 = re.match(r"(\d+)\.\s+(.+)", text)
            if m2:
                answers[int(m2.group(1))] = m2.group(2)

print("=== All Answers ===")
for k in sorted(answers.keys()):
    print(f"  Q{k}: {answers[k]}")

# Check distribution of letter answers
letter_answers = [v for v in answers.values() if len(str(v)) == 1 and str(v).upper() in "ABCDE"]
print(f"\n=== Answer Distribution (choice questions) ===")
print(f"Total: {len(letter_answers)}")
print(f"Distribution: {Counter(letter_answers)}")

# Check adjacent same answers
print(f"\n=== Adjacent Same Answer Check ===")
prev = None
for k in sorted(answers.keys()):
    v = answers[k]
    if len(str(v)) == 1 and str(v).upper() in "ABCDE":
        if v == prev:
            print(f"  WARNING: Q{k} has same answer as Q{k-1}: {v}")
        prev = v
    else:
        prev = None

# Check word bank (2 lines, 5 words each)
print(f"\n=== Word Bank Check ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "beach" in text and "delicious" in text and "terrible" in text:
        words = text.split()
        print(f"  Line 1: {len(words)} words: {text}")
    if "place" in text and "visit" in text and "ago" in text:
        words = text.split()
        print(f"  Line 2: {len(words)} words: {text}")

# Check five-choose-four options (A-E each on own line)
print(f"\n=== Five-Choose-Four Options Check ===")
w5_options = []
for p in doc.paragraphs:
    text = p.text.strip()
    if re.match(r"^[A-E]\.\s", text):
        w5_options.append(text)
        print(f"  {text}")

# Check cloze options (A/B/C on same line)
print(f"\n=== Cloze Options Check (should be same line) ===")
for p in doc.paragraphs:
    text = p.text.strip()
    m = re.match(r"^(\d+)\.\s+A\.\s+\S+\s+B\.\s+\S+\s+C\.\s+\S+", text)
    if m:
        print(f"  Q{m.group(1)}: OK - {text[:60]}...")

# Check for forbidden content
print(f"\n=== Forbidden Content Check ===")
irregular = ["went", "bought", "ate", "took", "came", "saw", "said", "made", "got", "gave", "ran", "sat", "wrote", "spoke", "drank", "sang", "swam", "drove", "flew", "slept", "fell", "told", "sold", "held", "found", "spent"]
full_text = "\n".join(p.text for p in doc.paragraphs)
# Only check cloze and grammar_fill sections (not reading comprehension)
cloze_section = False
gf_section = False
for p in doc.paragraphs:
    text = p.text.strip()
    if "完形填空" in text and "第一节" in text:
        cloze_section = True
    if "语法填空" in text and "第二节" in text:
        cloze_section = False
        gf_section = True
    if "综合技能" in text:
        gf_section = False
    
    if cloze_section or gf_section:
        words_lower = re.findall(r'\b[a-z]+\b', text.lower())
        for w in irregular:
            if w in words_lower:
                print(f"  WARNING: irregular verb '{w}' found in {'cloze' if cloze_section else 'grammar_fill'}: {text[:80]}")

print("\n=== Summary ===")
print(f"Total questions: {len(answers)}")
print(f"Expected: 51 (11 reading + 4 w5 + 10 cloze + 10 grammar_fill + 5 sa + 1 writing + 10 diag)")
