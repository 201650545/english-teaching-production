# -*- coding: utf-8 -*-
"""Check DOCX formatting details: fonts, sizes, indentation, spacing."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

path = r"D:\英语教学\邓兴华\第08课时\第08课时_配套练习.docx"
doc = Document(path)

print("=== Font Check ===")
# Check heading
for p in doc.paragraphs[:1]:
    for r in p.runs:
        print(f"Heading: font={r.font.name}, size={r.font.size}, bold={r.font.bold}")
        rpr = r._element.find(qn('w:rPr'))
        if rpr is not None:
            rf = rpr.find(qn('w:rFonts'))
            if rf is not None:
                print(f"  eastAsia={rf.get(qn('w:eastAsia'))}")

# Check a passage paragraph
for i, p in enumerate(doc.paragraphs):
    if "Yesterday was a busy" in p.text:
        for r in p.runs:
            print(f"\nPassage text: font={r.font.name}, size={r.font.size}, bold={r.font.bold}")
            rpr = r._element.find(qn('w:rPr'))
            if rpr is not None:
                rf = rpr.find(qn('w:rFonts'))
                if rf is not None:
                    print(f"  eastAsia={rf.get(qn('w:eastAsia'))}")
        # Check paragraph format
        pf = p.paragraph_format
        print(f"  space_after={pf.space_after}, space_before={pf.space_before}")
        # Check first line indent
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                print(f"  firstLineChars={ind.get(qn('w:firstLineChars'))}")
        break

# Check a question paragraph
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("1. When"):
        for r in p.runs:
            print(f"\nQuestion text: font={r.font.name}, size={r.font.size}, bold={r.font.bold}")
        pf = p.paragraph_format
        print(f"  space_after={pf.space_after}, space_before={pf.space_before}")
        break

# Check section heading
for i, p in enumerate(doc.paragraphs):
    if "第一部分" in p.text:
        for r in p.runs:
            print(f"\nSection heading: font={r.font.name}, size={r.font.size}, bold={r.font.bold}")
            rpr = r._element.find(qn('w:rPr'))
            if rpr is not None:
                rf = rpr.find(qn('w:rFonts'))
                if rf is not None:
                    print(f"  eastAsia={rf.get(qn('w:eastAsia'))}")
        break

# Check page count (rough estimate)
total_paras = len(doc.paragraphs)
print(f"\n=== Page Count Estimate ===")
print(f"Total paragraphs: {total_paras}")
# Rough: ~30 paragraphs per page
print(f"Estimated pages: {total_paras // 30} - {total_paras // 25}")

# Check for answer sheet
print(f"\n=== Answer Sheet Check ===")
found_answer_sheet = False
for p in doc.paragraphs:
    if "答题卡" in p.text:
        found_answer_sheet = True
        break
print(f"Answer sheet found: {found_answer_sheet}")

# Check for blueprint table (命题双向细目表)
found_blueprint = False
for p in doc.paragraphs:
    if "细目表" in p.text or "双向" in p.text:
        found_blueprint = True
        break
print(f"Blueprint table found: {found_blueprint}")

# Check tables in document
print(f"Tables in document: {len(doc.tables)}")
for i, t in enumerate(doc.tables):
    print(f"  Table {i}: {len(t.rows)} rows x {len(t.columns)} cols")
    if len(t.rows) > 0:
        print(f"    Header: {[c.text for c in t.rows[0].cells]}")
