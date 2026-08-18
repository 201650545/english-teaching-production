import os, re
try:
    from docx import Document
except ImportError:
    print('NO docx module'); raise SystemExit
base = r'D:\英语教学\邓兴华'
for lesson in [16,17,18,19,20]:
    fn = os.path.join(base, '第%02d课时' % lesson, '第%02d课时_配套练习_中等.docx' % lesson)
    if not os.path.exists(fn):
        print('L%d MISSING' % lesson); continue
    doc = Document(fn)
    txt = []
    for p in doc.paragraphs:
        txt.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                txt.append(c.text)
    full = '\n'.join(txt)
    past_kw = [' was ', ' were ', ' went ', ' came ', ' had ', ' did ', ' yesterday', ' last week',
               ' ago', ' cooked', ' played', ' visited', ' walked', ' watched', ' studied',
               ' bought', ' taught', ' ate ', ' drank ', ' broke', ' took', ' gave', ' made ',
               'Last year', 'Last weekend', 'two days ago', 'last month']
    hits = [k for k in past_kw if k in (' '+full+' ')]
    src_cnt = full.count('DXH2026_L') + full.count('source_id')
    print('L%d: source=%d past_hits=%s' % (lesson, src_cnt, hits if (lesson in (17,18,19) and hits) else 'NONE' if lesson in (17,18,19) else 'allowed'))
    print('   len=%d, sample head: %s' % (len(full), full[:80].replace('\n',' ')))