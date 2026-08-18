# -*- coding: utf-8 -*-
"""Verify the generated DOCX file."""
import os
from docx import Document

path = r"D:\英语教学\邓兴华\第08课时\第08课时_配套练习.docx"
print(f"File exists: {os.path.exists(path)}")
print(f"File size: {os.path.getsize(path)} bytes")

doc = Document(path)
print(f"Total paragraphs: {len(doc.paragraphs)}")
print()

# Print first 50 paragraphs to check structure
for i, p in enumerate(doc.paragraphs[:80]):
    text = p.text.strip()
    if text:
        print(f"[{i}] {text[:120]}")

print()
print("=== Checking key structure ===")

# Check for main heading
found_heading = False
for p in doc.paragraphs[:5]:
    if "第" in p.text and "课时" in p.text and "配套练习" in p.text:
        found_heading = True
        print(f"Main heading: {p.text}")
        break

# Check for sections
sections = []
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith("第一部分") or text.startswith("第二部分") or text.startswith("第三部分") or text.startswith("教学诊断"):
        sections.append(text)
print(f"Sections found: {sections}")

# Check for answer key
found_answers = False
for p in doc.paragraphs:
    if "参考答案" in p.text:
        found_answers = True
        break
print(f"Answer key found: {found_answers}")

# Count questions (look for numbered items)
question_count = 0
for p in doc.paragraphs:
    import re
    if re.match(r"^\d+\.", p.text.strip()):
        question_count += 1
print(f"Numbered items: {question_count}")

# Check for word bank format (each line 5 words)
for i, p in enumerate(doc.paragraphs):
    if "beach" in p.text and "delicious" in p.text:
        print(f"Word bank area [{i}]: {p.text}")
