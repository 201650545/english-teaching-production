# -*- coding: utf-8 -*-
"""配套练习生成器（真实题目版 · exam_spec v2026.2 结构）v2
重设计（2026-08-02 教师反馈逐条落地）：
1. 语篇分段 + 首行缩进 1 字符
2. 选择题选项竖向对齐（A/B/C 各占一行同列）
3. 五选四 / 选词填空：答案直接填入文中空内，不单列小题
4. 完形填空：语篇分段缩进；选项竖向对齐
5. 语法复盘：5 选择题 + 5 填空题
6. 题号全卷连续：内容文件空号小节内 1..N，本器按部分全局重编号
7. 答案选项随机化：seed=课时号 打乱选项，相邻两题正确字母不得相同（跨题型跟踪）
数据来源：本课内容文件 practice_content_{学生}_{课}.json（教师可审）。
难度由 tier 决定（邓兴华=中等）。卷末附参考答案 + 双向细目表 + 溯源登记。
"""
import json, os, re, random
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
QI = Cm(0.6)    # 题号/正文缩进
OPT = Cm(1.5)   # 选项缩进（A/B/C 同列竖向对齐）
BCOL = Cm(7.0)  # B 列制表位（从左页边距起算，对齐 2026 中考固定列 x≈221pt）
CCOL = Cm(13.0) # C 列制表位（对齐 2026 中考固定列 x≈388pt）

def _set_font(run, cn="宋体", en="Times New Roman", size=12, bold=False, color=None):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), cn)
    if color:
        run.font.color.rgb = RGBColor(*color)

def _para(doc, text, size=10.5, bold=False, align=None, cn="宋体", en="Times New Roman",
          space_after=2, left_indent=None, first_indent_chars=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if first_indent_chars:
        ind = p._p.get_or_add_pPr().get_or_add_ind()
        ind.set(qn("w:firstLineChars"), str(first_indent_chars * 100))
        ind.set(qn("w:firstLine"), "0")
    run = p.add_run(text)
    _set_font(run, cn=cn, en=en, size=size, bold=bold)
    return p

def _heading(doc, text, size=16, color=(230, 57, 70)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    _set_font(run, cn="黑体", en="Arial", size=size, bold=True, color=color)
    return p

def _section(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run, cn="黑体", en="Arial", size=size, bold=True, color=(230, 57, 70))
    return p

def _sub(doc, text, size=10.5, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    _set_font(run, cn="黑体", en="Arial", size=size, bold=bold, color=(40, 40, 40))
    return p

def _passage(doc, paragraphs, first_indent=1):
    """语篇：分段渲染，每段首行缩进 first_indent 字符（规范 02 §2.2：首行缩进 1 字符）。
    支持段落中含 \n 的多行文本（对话体）。"""
    if isinstance(paragraphs, str):
        paragraphs = [paragraphs]
    for para in paragraphs:
        # 对话体（含换行）每行独立渲染，不缩进对话行
        if "\n" in para:
            for line in para.split("\n"):
                line = line.strip()
                if line:
                    _para(doc, line, size=10.5, space_after=2, first_indent_chars=0)
        else:
            _para(doc, para, size=10.5, space_after=3, first_indent_chars=first_indent)

def _renumber(paragraphs, start):
    """把语篇中的空号 ___N___ 从小节内编号改为全局编号（偏移 start-1）。"""
    def repl(m):
        return "___%d___" % (int(m.group(1)) + start - 1)
    return [re.sub(r"___(\d+)___", repl, p) for p in paragraphs]

# ───────── 答案选项随机化（2026-08-03 教师规范：连续两题不得同字母，尽可能随机） ─────────
def _shuffle_one(q, rng, prev):
    """随机重排单题 A/B/C 选项（opts 形如 [['A',文本],...]），避免与上一题正确字母重复；返回本题意向字母。"""
    pairs = [list(p) for p in q["opts"]]
    n = len(pairs)
    letters = "ABCDEFGH"
    cur = letters.index(q["answer"].upper()) % n
    order = list(range(n)); rng.shuffle(order)
    def letter_at(i):
        return letters[order.index(i)]
    for _ in range(30):
        if prev is None or letter_at(cur) != prev:
            break
        alt = (cur + 1) % n
        ci, ai = order.index(cur), order.index(alt)
        order[ci], order[ai] = order[ai], order[ci]
    q["opts"] = [[letters[i], pairs[order[i]][1]] for i in range(n)]
    q["answer"] = letter_at(cur)
    return q["answer"]

def _rand_w5(w, rng, prev):
    """五选四：随机分配 A-E 候选字母，空1 与上一题不重复；返回最后一空的字母。"""
    pairs = list(w["candidates"])
    rng.shuffle(pairs)
    letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"[:len(pairs)]
    old_to_new = {}
    for i, (old, _s) in enumerate(pairs):
        old_to_new[old] = letters[i]
    w["candidates"] = [[letters[i], s] for i, (_old, s) in enumerate(pairs)]
    w["answers"] = {k: old_to_new[v] for k, v in w["answers"].items()}
    bl = sorted(w["answers"], key=int)
    while w["answers"][bl[0]] == prev:
        m = {old: letters[(i + 1) % len(letters)] for i, old in enumerate(letters)}
        w["answers"] = {k: m[v] for k, v in w["answers"].items()}
        for cand in w["candidates"]:
            cand[0] = m[cand[0]]
    return w["answers"][bl[-1]]

def randomize_all(content, seed):
    """按卷面顺序统一打乱所有选择题选项，保证相邻两题正确字母不同（seed=lesson 可复现）。"""
    rng = random.Random(seed)
    prev = None
    for tag in ("a", "b", "c"):
        for q in content["reading_%s" % tag]["questions"]:
            prev = _shuffle_one(q, rng, prev)
    prev = _rand_w5(content["w5"], rng, prev)
    for it in content["cloze"]["items"]:
        prev = _shuffle_one(it, rng, prev)
    for q in content["grammar_diag"]["mc"]:
        prev = _shuffle_one(q, rng, prev)

def _question(doc, num, stem, size=10.5):
    _para(doc, "%d. %s" % (num, stem), size=size, left_indent=QI, space_after=2)

def _options3(doc, opts, num=None, size=10.5):
    """中考三列选项排版（2026 湖南中考实测：A≈x64 / B≈x221 / C≈x399 固定列）。
    num 给定时同排显示题号（完形用），否则从选项列开始（阅读/语法选择用）。
    B/C 用绝对制表位（从左页边距起算 7.0/13.0cm），A→B≈5.5cm、B→C≈6.0cm，
    全卷选项共用同一竖列，选项文本不得长于列宽（内容侧须控制长度）。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6) if num else OPT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(BCOL)
    p.paragraph_format.tab_stops.add_tab_stop(CCOL)
    prefix = ("%d. " % num) if num else ""
    text = prefix + "\t".join("%s. %s" % (l, t) for l, t in opts)
    run = p.add_run(text)
    _set_font(run, cn="宋体", en="Times New Roman", size=size)

def _ans_runs(answers, max_per_line=5):
    """中考答案页紧凑排列：连续字母题连排为 `1~5 CCBBA`（每行≤5个），词/句题逐行 `N. 答案`。"""
    lines, i, n = [], 0, len(answers)
    while i < n:
        num, a = answers[i]
        if len(a) == 1 and a.upper() in "ABCDE":
            j = i
            while j < n and len(answers[j][1]) == 1 and answers[j][1].upper() in "ABCDE":
                j += 1
            for k in range(i, j, max_per_line):
                chunk = answers[k:k + max_per_line]
                if len(chunk) == 1:
                    lines.append("%d. %s" % (chunk[0][0], chunk[0][1]))
                else:
                    lines.append("%d~%d %s" % (chunk[0][0], chunk[-1][0], "".join(x[1] for x in chunk)))
            i = j
        else:
            lines.append("%d. %s" % (num, a))
            i += 1
    return lines

def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, size=9, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            _set_font(run, size=9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def _render_answer_card(doc, ans_list):
    """渲染答题卡填涂格（选择题）。"""
    # 每行4题，每题显示 [A] [B] [C]
    i = 0
    while i < len(ans_list):
        chunk = ans_list[i:i+4]
        parts = []
        for n, _ in chunk:
            parts.append("%d. [A] [B] [C]" % n)
        _para(doc, "  ".join(parts), size=10.5, space_after=1)
        i += 4

# ─────────────────────────── 题型渲染 ───────────────────────────
def _render_reading_choice(doc, content, qnum):
    ans, bp = [], []
    for tag in ("a", "b", "c"):
        pg = content["reading_%s" % tag]
        _sub(doc, "Passage %s（%s · %s · %d 词）" % (tag.upper(), pg.get("genre", ""), pg.get("difficulty", ""), pg.get("word_count", "")))
        _passage(doc, pg["paragraphs"])
        for q in pg["questions"]:
            _question(doc, qnum, q["q"])
            _options3(doc, q["opts"])
            ans.append((qnum, q["answer"]))
            bp.append({"题号": qnum, "题型": "阅读选择·%s篇" % tag.upper(), "分值": 2,
                       "绑定考点/词": pg.get("绑定", "运动·喜好表达"), "难度": "中", "母本源ID": pg["id"]})
            qnum += 1
    return qnum, ans, bp

def _render_w5(doc, content, qnum):
    w = content["w5"]
    ans, bp = [], []
    if w.get("title"):
        _sub(doc, w["title"])
    _passage(doc, _renumber(w["paragraphs"], qnum))
    _para(doc, "方框：", size=10.5, bold=True, space_after=1)
    for letter, s in w["candidates"]:
        _para(doc, "%s. %s" % (letter, s), size=10.5, left_indent=OPT, space_after=1)
    for k, v in w["answers"].items():
        g = qnum + int(k) - 1
        ans.append((g, v))
        bp.append({"题号": g, "题型": "五选四", "分值": 2, "绑定考点/词": "运动建议", "难度": "中", "母本源ID": w["id"]})
    return qnum + len(w["answers"]), ans, bp

def _render_cloze(doc, content, qnum):
    c = content["cloze"]
    ans, bp = [], []
    _sub(doc, c.get("title", "阅读短文，从每题所给的 A、B、C 三个选项中选出最佳选项"))
    _passage(doc, _renumber(c["paragraphs"], qnum))
    for it in c["items"]:
        g = qnum + it["num"] - 1
        _options3(doc, it["opts"], num=g)
        ans.append((g, it["answer"]))
        bp.append({"题号": g, "题型": "完形填空", "分值": 1.5, "绑定考点/词": c.get("绑定", "祈使句/What/like"),
                   "难度": "中", "母本源ID": c["id"]})
    return qnum + len(c["items"]), ans, bp

def _render_grammar_fill(doc, content, qnum):
    """语法填空（对齐 2026 湖南中考语言运用第二节）：词框 + 语篇含空（所给词变形填空），不单列小题。
    词框按每行5个排列（共两行），对齐规范 02 §7.2。"""
    wb = content["grammar_fill"]
    ans, bp = [], []
    _sub(doc, wb.get("title", "从方框内选择适当的词并用其正确形式填空"))
    # 词框：每行5个，共两行
    words = wb["word_bank"]
    half = (len(words) + 1) // 2
    for i in range(0, len(words), half):
        _para(doc, "  ".join(words[i:i+half]), size=10.5, bold=True, space_after=1)
    # 语篇分段：对话按换行分组渲染
    paras = wb["paragraphs"]
    if isinstance(paras, str):
        paras = [paras]
    _passage(doc, _renumber(paras, qnum))
    for i, a in enumerate(wb["answers"]):
        g = qnum + i
        ans.append((g, a))
        bp.append({"题号": g, "题型": "语法填空", "分值": 1, "绑定考点/词": wb.get("绑定", "G19·G20·G21"),
                   "难度": "中", "母本源ID": wb["id"]})
    return qnum + len(wb["answers"]), ans, bp

def _render_sa(doc, content, qnum):
    sa = content["sa"]
    ans, bp = [], []
    _sub(doc, sa.get("title", "阅读短文，回答下列问题或按要求完成句子"))
    _passage(doc, sa["paragraphs"])
    for i, q in enumerate(sa["questions"]):
        n = qnum + i
        _question(doc, n, q["q"])
        _para(doc, "    ____________________________________________________", size=10.5)
        ans.append((n, q["answer"]))
        bp.append({"题号": n, "题型": q.get("type", "简答"), "分值": 2, "绑定考点/词": sa.get("绑定", "运动/祈使句"),
                   "难度": "中", "母本源ID": sa["id"]})
    return qnum + len(sa["questions"]), ans, bp

def _render_writing(doc, content, qnum):
    wr = content["writing"]
    if wr.get("title"):
        _sub(doc, wr["title"])
    _para(doc, "%d. %s" % (qnum, wr["prompt"]), size=10.5, left_indent=QI, space_after=2)
    _para(doc, wr["requirements"], size=10.5, left_indent=QI, space_after=4)
    for _ in range(6):
        _para(doc, "    ________________________________________________________", size=10.5, space_after=2)
    bp = [{"题号": qnum, "题型": "书面表达", "分值": 15, "绑定考点/词": wr.get("绑定", "运动·写作"), "难度": "中", "母本源ID": wr["id"]}]
    return qnum + 1, [(qnum, "见参考答案范文")], bp

def _render_grammar(doc, content, qnum):
    g = content["grammar_diag"]
    ans, bp = [], []
    _sub(doc, g.get("title", "语法复盘"))
    _sub(doc, "（一）单项选择（共 5 小题，每小题 2 分）")
    for q in g["mc"]:
        _question(doc, qnum, q["q"])
        _options3(doc, q["opts"])
        ans.append((qnum, q["answer"]))
        bp.append({"题号": qnum, "题型": "语法诊断·选择", "分值": 2, "绑定考点/词": "G13·G14·G15",
                   "难度": "中", "母本源ID": g["id"]})
        qnum += 1
    _sub(doc, "（二）根据句意，用括号中所给词的适当形式填空（共 5 小题，每小题 2 分）")
    for q in g["fill"]:
        _question(doc, qnum, q["q"])
        ans.append((qnum, q["answer"]))
        bp.append({"题号": qnum, "题型": "语法诊断·填空", "分值": 2, "绑定考点/词": "G13·G15",
                   "难度": "中", "母本源ID": g["id"]})
        qnum += 1
    return qnum, ans, bp

# ─────────────────────────── 主构建 ───────────────────────────
def build_practice(card, content, out_path):
    lesson, student, tier = card["lesson"], card["student"], card["tier"]
    theme = card["theme"]
    # 答案选项随机化 + 连续两题不同字母（2026-08-03 教师规范）
    randomize_all(content, lesson)
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)

    _heading(doc, "第 %02d 课时配套练习" % lesson)
    _para(doc, "学生：%s    层级：%s    主题：%s    难度：中等" % (student, tier, theme),
          align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    _para(doc, "姓名：____________    得分：____________    用时：____________",
          align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    _para(doc, "结构对齐 2026 湖南中考（不含听力）· 满分：80 分（含听力 100 分）", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    qnum = 1
    ans1_rd, ans1_w5 = [], []   # 阅读选择 / 五选四
    ans2_cl, ans2_gf = [], []   # 完形 / 语法填空
    ans3_sa, ans3_wr = [], []   # 简答 / 书面表达
    ans4_mc, ans4_fl = [], []   # 诊断选择 / 诊断填空

    # 第一部分 阅读理解（对齐中考 第二部分）
    _section(doc, "第一部分　阅读理解（共两节，满分 30 分）")
    _sub(doc, "第一节（共 11 小题，每小题 2 分）——阅读下列材料，从每题所给的 A、B、C 三个选项中选出最佳选项")
    qnum, ans1_rd, _ = _render_reading_choice(doc, content, qnum)
    _sub(doc, "第二节（共 4 小题，每小题 2 分）——阅读下面短文，从方框中选出可以填入空白处的最佳选项（有一项为多余选项）")
    qnum, ans1_w5, _ = _render_w5(doc, content, qnum)

    # 第二部分 语言运用（对齐中考 第三部分）
    _section(doc, "第二部分　语言运用（共两节，满分 25 分）")
    _sub(doc, "第一节　完形填空（共 10 小题，每小题 1.5 分）")
    qnum, ans2_cl, _ = _render_cloze(doc, content, qnum)
    _sub(doc, "第二节　语法填空（共 10 小题，每小题 1 分）")
    qnum, ans2_gf, _ = _render_grammar_fill(doc, content, qnum)

    # 第三部分 综合技能（对齐中考 第四部分）
    _section(doc, "第三部分　综合技能（共两节，满分 25 分）")
    _sub(doc, "第一节　阅读表达（共 5 小题，每小题 2 分）")
    qnum, ans3_sa, _ = _render_sa(doc, content, qnum)
    _sub(doc, "第二节　书面表达（满分 15 分）")
    qnum, ans3_wr, _ = _render_writing(doc, content, qnum)

    # 教学诊断附件（教师课后诊断用，不计入考试；非中考结构，独立标注）
    _section(doc, "教学诊断附件（教师课后诊断用 · 不计入考试）")
    _sub(doc, "语法复盘（5 个单项选择 + 5 个填空）")
    qnum, diag_ans, _ = _render_grammar(doc, content, qnum)
    n_mc = len(content["grammar_diag"]["mc"])
    ans4_mc, ans4_fl = diag_ans[:n_mc], diag_ans[n_mc:]

    # ── 参考答案（对齐 2026 中考答案页：字母题连排、词句题逐行） ──
    doc.add_page_break()
    _section(doc, "参考答案")
    _sub(doc, "第一部分　阅读理解")
    for ln in _ans_runs(ans1_rd):
        _para(doc, ln, size=10.5, space_after=1)
    for ln in _ans_runs(ans1_w5):
        _para(doc, ln, size=10.5, space_after=1)
    _sub(doc, "第二部分　语言运用")
    _sub(doc, "第一节　完形填空")
    for ln in _ans_runs(ans2_cl):
        _para(doc, ln, size=10.5, space_after=1)
    _sub(doc, "第二节　语法填空")
    for n, a in ans2_gf:
        _para(doc, "%d. %s" % (n, a), size=10.5, space_after=1)
    _sub(doc, "第三部分　综合技能")
    _sub(doc, "第一节　阅读表达")
    for n, a in ans3_sa:
        _para(doc, "%d. %s" % (n, a), size=10.5, space_after=1)
    _sub(doc, "第二节　书面表达")
    _passage(doc, content["writing"]["sample"])
    _sub(doc, "教学诊断附件")
    for ln in _ans_runs(ans4_mc):
        _para(doc, ln, size=10.5, space_after=1)
    for n, a in ans4_fl:
        _para(doc, "%d. %s" % (n, a), size=10.5, space_after=1)

    # ── 答题卡 ──
    doc.add_page_break()
    _section(doc, "答题卡")
    _sub(doc, "第一部分　阅读理解（用 2B 铅笔填涂）")
    # 阅读选择 + 五选四
    _render_answer_card(doc, ans1_rd + ans1_w5)
    _sub(doc, "第二部分　语言运用")
    _sub(doc, "第一节　完形填空")
    _render_answer_card(doc, ans2_cl)
    _sub(doc, "第二节　语法填空（直接填写单词）")
    for n, a in ans2_gf:
        _para(doc, "%d. ________" % n, size=10.5, space_after=1)
    _sub(doc, "第三部分　综合技能")
    _sub(doc, "第一节　阅读表达（在横线上作答）")
    for n, a in ans3_sa:
        _para(doc, "%d. ____________________________________________" % n, size=10.5, space_after=1)
    _sub(doc, "第二节　书面表达（在划线区域作答）")
    for _ in range(6):
        _para(doc, "________________________________________________________", size=10.5, space_after=2)
    _sub(doc, "教学诊断附件")
    _render_answer_card(doc, ans4_mc)
    _sub(doc, "填空题")
    for n, a in ans4_fl:
        _para(doc, "%d. ________" % n, size=10.5, space_after=1)

    # ── 命题双向细目表 ──
    doc.add_page_break()
    _section(doc, "命题双向细目表")
    bp_rows = []
    # 阅读选择 (q1-11)
    qn = 1
    for tag in ("a", "b", "c"):
        pg = content["reading_%s" % tag]
        for i, q in enumerate(pg["questions"]):
            bp_rows.append([str(qn), "阅读选择·%s篇" % tag.upper(), str(2), pg.get("绑定", ""), "中", pg["id"]])
            qn += 1
    # 五选四 (q12-15)
    w = content["w5"]
    for k in sorted(w["answers"], key=int):
        bp_rows.append([str(qn), "五选四", str(2), w.get("绑定", ""), "中", w["id"]])
        qn += 1
    # 完形 (q16-25)
    c = content["cloze"]
    for it in c["items"]:
        bp_rows.append([str(qn), "完形填空", str(1.5), c.get("绑定", ""), "中", c["id"]])
        qn += 1
    # 语法填空 (q26-35)
    wb = content["grammar_fill"]
    for i in range(len(wb["answers"])):
        bp_rows.append([str(qn), "语法填空", str(1), wb.get("绑定", ""), "中", wb["id"]])
        qn += 1
    # 阅读表达 (q36-40)
    sa = content["sa"]
    for i, q in enumerate(sa["questions"]):
        bp_rows.append([str(qn), q.get("type", "简答"), str(2), sa.get("绑定", ""), "中", sa["id"]])
        qn += 1
    # 书面表达 (q41)
    wr = content["writing"]
    bp_rows.append([str(qn), "书面表达", str(15), wr.get("绑定", ""), "中", wr["id"]])
    qn += 1
    # 语法诊断 (q42-51)
    gd = content["grammar_diag"]
    for q in gd["mc"]:
        bp_rows.append([str(qn), "语法诊断·选择", str(2), q.get("绑定", ""), "中", gd["id"]])
        qn += 1
    for i, q in enumerate(gd["fill"]):
        bp_rows.append([str(qn), "语法诊断·填空", str(2), q.get("绑定", ""), "中", gd["id"]])
        qn += 1
    _table(doc, ["题号", "题型", "分值", "绑定考点/词", "难度", "母本源ID"], bp_rows)

    doc.save(out_path)
    return out_path

if __name__ == "__main__":
    card = {
        "lesson": 5, "student": "邓兴华", "tier": "中等", "stage": "S2", "type": "normal",
        "grammar": ["Do/Don't 祈使句", "What 特殊疑问句", "like + 名词 / like to do"],
        "theme": "运动与喜好表达", "vocab": {"new_count": 20, "review_count": 0, "theme": "sports"},
        "phonics": "bl/cl/fl/gl/pl/sl", "listening": False,
    }
    content = json.load(open(os.path.join(HERE, "practice_content_DXH_L05.json"), encoding="utf-8"))
    out = os.path.join(os.path.dirname(HERE), "邓兴华", "第05课", "第05课_配套练习_中等.docx")
    p = build_practice(card, content, out)
    print("配套练习生成：%s (%d bytes)" % (p, os.path.getsize(p)))
